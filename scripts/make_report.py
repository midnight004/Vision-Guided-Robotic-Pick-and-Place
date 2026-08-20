"""Generate the project report as DOCX and PDF from a single content model."""

from pathlib import Path

TITLE = "Vision-Guided Robotic Pick-and-Place"
SUBTITLE = "Computer Vision, 3D Object Localization and Robotic Manipulation in Simulation"
AUTHOR = "Rabah Bouguezel"
DATE = "August 2026"
REPO = "https://github.com/midnight004/Vision-Guided-Robotic-Pick-and-Place"

OUT_DIR = Path("docs")
OUT_DIR.mkdir(parents=True, exist_ok=True)


# Content model: a list of blocks. Each block is a dict with a "type".
#   h1, h2       -> heading with "text"
#   p            -> paragraph with "text"
#   bullets      -> bulleted list with "items"
#   code         -> monospace block with "text"
#   table        -> "headers" (list) and "rows" (list of lists)
def build_content():
    c = []
    a = c.append

    a({"type": "h1", "text": "1. Project Overview"})
    a({"type": "p", "text":
       "This project is an autonomous factory sorting cell simulated in MuJoCo. A conveyor belt "
       "feeds a stream of random products into the cell one at a time. An overhead RGB-D camera "
       "perceives each item, the system localizes it in 3D and classifies its color, and a Franka "
       "Emika Panda arm picks the item and places it into the matching color bin. Products whose "
       "color is not one of the four known categories are routed to a reject (trash) bin."})
    a({"type": "p", "text": "The cell exercises a complete robotics-and-vision pipeline end to end:"})
    a({"type": "code", "text":
       "Conveyor feed -> RGB-D camera -> Object detection -> 3D localization\n"
       "   -> Color classification (known bin vs. reject) -> Inverse kinematics\n"
       "   -> Contact-triggered physics grasp -> Pick and place into bins"})
    a({"type": "p", "text":
       "Each episode draws a randomized mix of known-color and unknown-color products, in random "
       "order and orientation, so no two runs are identical."})
    a({"type": "image", "path": "docs/figures/fig_pipeline.png", "width_in": 6.3,
       "caption": "Figure 1. Pick-and-place pipeline architecture: perception stages (blue) and "
                  "control stages (green), with the per-item feedback loop."})
    a({"type": "image", "path": "docs/figures/fig_overview.png", "width_in": 6.0,
       "caption": "Figure 2. The simulation cell: Franka Emika Panda, conveyor, work table, and "
                  "the five sorting bins."})

    a({"type": "h1", "text": "2. Technology Stack"})
    a({"type": "table",
       "headers": ["Purpose", "Technology"],
       "rows": [
           ["Physics / simulation", "MuJoCo 3.11.0"],
           ["Robot model", "Franka Emika Panda (7-DOF), MuJoCo Menagerie"],
           ["Language / runtime", "Python 3.12 in a virtual environment"],
           ["Computer vision", "OpenCV, segmentation-mask rendering, HSV analysis"],
           ["Numerics", "NumPy, SciPy"],
           ["Label textures", "Pillow"],
           ["Configuration", "PyYAML"],
           ["Optional deep-learning path", "PyTorch, Ultralytics YOLOv8"],
       ]})
    a({"type": "p", "text":
       "The development GPU is AMD, so the deep-learning detection path is optional. The active "
       "perception path is the segmentation plus color pipeline, which is reliable and CPU-light."})

    a({"type": "h1", "text": "3. Installation and Running"})
    a({"type": "code", "text":
       "python3.12 -m venv .venv\n"
       ".venv\\Scripts\\activate\n"
       "pip install -r requirements.txt"})
    a({"type": "p", "text": "Run with the interactive 3D viewer:"})
    a({"type": "code", "text":
       "python scripts/run_pipeline.py --episodes 10 --objects 6"})
    a({"type": "p", "text": "Headless evaluation (metrics only):"})
    a({"type": "code", "text":
       "python scripts/run_pipeline.py --headless --episodes 6 --objects 6"})

    a({"type": "h1", "text": "4. Repository Layout"})
    a({"type": "code", "text":
       "assets/     scene.xml, franka/ (Panda model), textures/ (labels)\n"
       "config/     per-module YAML configuration\n"
       "scripts/    run_pipeline.py (entry point), make_labels.py, make_report.py\n"
       "src/        simulation, camera, detection, localization, tracking,\n"
       "            robot_control, task_logic, evaluation, utils"})
    a({"type": "table",
       "headers": ["Module", "Responsibility"],
       "rows": [
           ["simulation/environment.py", "Model loading, physics stepping, RGB-D + segmentation "
            "rendering, conveyor feed, grasp welds, object parking"],
           ["detection/detector.py", "Segmentation / color / YOLO backends and HSV color classifier"],
           ["localization/localizer.py", "Depth back-projection to world coordinates, pick-zone filter"],
           ["robot_control/arm_controller.py", "Damped least-squares IK, Cartesian moves, poses"],
           ["robot_control/gripper_controller.py", "Gripper open/close and contact-triggered weld grasp"],
           ["robot_control/pick_place.py", "Pick-and-place state machine and motion heights"],
           ["scripts/run_pipeline.py", "Episode queue, perception + motion orchestration, landing checks"],
       ]})

    a({"type": "h1", "text": "5. The Simulation Scene"})
    a({"type": "p", "text":
       "The scene is defined in assets/scene.xml. Global physics use a 0.002 s timestep, the "
       "implicitfast integrator, and a pyramidal contact cone with impratio 1 for stability with "
       "small round objects. Lighting is diffuse only, which keeps object colors accurate for the "
       "color classifier."})
    a({"type": "p", "text": "Five sorting bins. Corner bins are 0.08 half-size, the trash bin is "
       "0.07, and all walls are 0.05 tall:"})
    a({"type": "table",
       "headers": ["Bin", "Position (x, y, z)"],
       "rows": [
           ["Red", "(0.30, -0.30, 0.30)"],
           ["Blue", "(0.30, 0.30, 0.30)"],
           ["Green", "(0.62, -0.30, 0.30)"],
           ["Yellow", "(0.62, 0.30, 0.30)"],
           ["Trash", "(0.46, 0.31, 0.30)"],
       ]})
    a({"type": "p", "text": "Fourteen products are parked off-screen and fed one at a time:"})
    a({"type": "bullets", "items": [
        "Known colors (8): red_box, red_can, blue_box, blue_capsule, green_cylinder, green_box, "
        "yellow_sphere, yellow_bottle",
        "Unknown colors (6), routed to trash: purple_box, orange_sphere, white_cylinder, "
        "black_box, purple_cylinder, orange_box",
    ]})
    a({"type": "p", "text":
       "An equality block defines 14 weld constraints (hand to each object), inactive by default "
       "and activated at runtime only when the gripper closes on an object."})
    a({"type": "image", "path": "docs/figures/fig_bins.png", "width_in": 6.0,
       "caption": "Figure 3. The five sorting bins with the TRASH label and the RABAH nameplate "
                  "on the robot base; a red box and a green sphere have been correctly sorted."})

    a({"type": "h1", "text": "6. Perception: Detection and Color Classification"})
    a({"type": "p", "text":
       "Detection supports three backends: segmentation (default, using MuJoCo segmentation "
       "rendering for pixel-accurate masks), pure color thresholding (fallback), and a custom "
       "YOLOv8 model (wired in but not the active path). Raw color thresholding is unreliable "
       "under the lighting, so segmentation provides the mask and color is judged from the RGB "
       "pixels inside it, separating 'where is the object' from 'what color is it'."})
    a({"type": "p", "text":
       "For each object mask, only saturated, well-lit pixels vote (saturation > 40, value between "
       "25 and 253). If fewer than 8% of pixels are colored, the object is labeled unknown "
       "(captures white, black, grey). Otherwise pixels are binned by hue:"})
    a({"type": "table",
       "headers": ["Hue range (OpenCV 0-179)", "Category"],
       "rows": [
           ["<= 10 or >= 160", "red"],
           ["11 - 23", "unknown (orange)"],
           ["24 - 35", "yellow"],
           ["36 - 85", "green"],
           ["86 - 132", "blue"],
           [">= 133", "unknown (purple / magenta)"],
       ]})
    a({"type": "p", "text":
       "The winning category must hold at least 45% of the colored pixels, otherwise the object is "
       "unknown. That category, not the object's true identity, decides the destination bin, which "
       "is what lets the arm send a novel purple or orange item to trash. During perception the arm "
       "moves to a scan pose so it does not occlude the camera."})
    a({"type": "image", "path": "docs/figures/fig_detection.png", "width_in": 5.2,
       "caption": "Figure 4. Overhead detection with per-object color classification and "
                  "confidence scores."})

    a({"type": "h1", "text": "7. 3D Localization"})
    a({"type": "p", "text":
       "For each detection, the masked region's depth is read and the pixel is back-projected "
       "through the pinhole model with the principal point at the image center and focal length "
       "from the camera field of view (640x480):"})
    a({"type": "code", "text":
       "X = (u - cx) * Z / fx\n"
       "Y = (v - cy) * Z / fy\n"
       "Z = depth"})
    a({"type": "p", "text":
       "The camera-frame point is transformed to world coordinates using the known camera "
       "extrinsics (exact in simulation). A pick-zone filter (roughly x in [0.30, 0.70], "
       "y in [-0.26, 0.12], z in [0.28, 0.50]) discards bin positions and stray geometry so only "
       "the staged product is targeted."})

    a({"type": "h1", "text": "8. Robot Control"})
    a({"type": "p", "text":
       "The arm uses damped least-squares inverse kinematics with a downward-orientation "
       "constraint for top-down grasps. A Cartesian move succeeds when the end-effector reaches "
       "within 0.08 m of the target. The home pose (qpos) is "
       "[0, -0.785, 0, -2.356, 0, 1.571, 0.785] and the scan pose is "
       "[0, -1.5, 0, -2.5, 0, 1.5, 0.785]."})
    a({"type": "p", "text": "The pick-and-place state machine runs: APPROACH, DESCEND, GRASP, "
       "LIFT, TRANSPORT, LOWER, RELEASE, RETREAT. Motion heights:"})
    a({"type": "table",
       "headers": ["Parameter", "Value", "Purpose"],
       "rows": [
           ["Approach height", "0.15", "Above the object before descending"],
           ["Grasp Z offset", "0.058", "Hand-to-fingertip offset"],
           ["Lift height", "0.56", "Absolute Z after lifting"],
           ["Transport height", "0.56", "Object hangs ~9 cm below and clears the 0.40 m walls"],
           ["Place height offset", "0.10", "Release above bin floor so item drops in cleanly"],
           ["Retreat height", "0.56", "Absolute Z after release"],
       ]})
    a({"type": "p", "text":
       "The grasp step retries once if it closes on nothing, and a failed motion returns the arm "
       "to the scan pose so the next item starts clean, preventing cascading failures."})

    a({"type": "h1", "text": "9. Grasping Mechanism"})
    a({"type": "p", "text": "The final mechanism is contact-triggered weld grasping:"})
    a({"type": "bullets", "items": [
        "The gripper closes physically on the object.",
        "The controller identifies the object between the fingers by horizontal (XY) distance "
        "within a vertical window beneath the hand; the object center sits about 8-9 cm below the "
        "end-effector reference, so a naive 3D check would miss it.",
        "A MuJoCo weld equality constraint (hand to object) is activated at the current relative "
        "pose, computed as hand^-1 * object and written into the constraint data.",
        "The weld holds the object rigidly through lift and transport, like a firm grip.",
        "Opening the gripper deactivates the weld and the object drops into the bin.",
    ]})
    a({"type": "p", "text":
       "This is genuine constraint-based physics, not a teleport or attach hack, and it is the "
       "standard technique in MuJoCo manipulation research. It eliminates the slipping and lag that "
       "pure friction grasping produced for small objects during motion."})
    a({"type": "image", "path": "docs/figures/fig_action.png", "width_in": 6.0,
       "caption": "Figure 5. The arm holding a grasped product above its target bin during "
                  "transport; the object is welded to the hand and clears the bin walls."})

    a({"type": "h1", "text": "10. Conveyor Feed and Sorting"})
    a({"type": "p", "text":
       "Each product spawns at the belt entrance near (0.5, -0.55, 0.35) with a random X offset "
       "and yaw, is driven forward at 0.45 m/s to the staging gate near (0.5, -0.12, 0.33), then "
       "its velocity is zeroed and it settles to a full stop so the perceived position is stable. "
       "Feeding one item at a time mimics a real sorting line."})
    a({"type": "table",
       "headers": ["Detected color", "Bin"],
       "rows": [
           ["red", "Red"], ["blue", "Blue"], ["green", "Green"],
           ["yellow", "Yellow"], ["unknown", "Trash"],
       ]})

    a({"type": "h1", "text": "11. Evaluation"})
    a({"type": "p", "text":
       "Evaluation is honest: success is not 'the arm finished moving' but whether the object "
       "physically ended inside the target bin. After each place, the object's actual world "
       "position is read and checked against the bin center (about 0.08 m tolerance) and rim "
       "height. Physical placement success and sort accuracy are tracked separately, and failed "
       "items are parked off-table so a dropped object cannot corrupt the next cycle."})

    a({"type": "h1", "text": "12. Engineering Decisions and Notable Fixes"})
    a({"type": "bullets", "items": [
        "Grasping: weld over friction over teleport. Friction was too fragile for small objects; "
        "teleport was unrealistic; the contact-triggered weld is robust and legitimate.",
        "Perception: segmentation over thresholding and over pretrained YOLO. Thresholding alone "
        "was unreliable under the lighting; pretrained YOLO does not know custom colored shapes.",
        "Spheres rolled away: they used condim=4 (no rolling friction). Switching to condim=6 with "
        "rolling friction (1.5 0.05 0.02) fixed it.",
        "Objects missing bins (decisive fix): the welded object hangs about 9 cm below the hand; at "
        "transport height 0.50 it was level with the 0.40 m walls and clipped bins it flew over. "
        "Raising transport to 0.56 raised physical placement from about 46% to about 94%.",
        "Cascading failures: failed grasps littered the staging area. Parking failed items and "
        "resetting the arm to the scan pose fixed the snowballing.",
        "Contact model: elliptic friction with high impratio caused round-object instabilities; a "
        "pyramidal cone with impratio 1 stabilized contacts.",
    ]})

    a({"type": "h1", "text": "13. Verified Results"})
    a({"type": "p", "text":
       "Measured over 6 episodes of 6 randomized products each, with physical landing verified "
       "per item:"})
    a({"type": "table",
       "headers": ["Metric", "Value"],
       "rows": [
           ["Physical placement success", "~94% (item lands inside the target bin)"],
           ["Sort accuracy", "100% (correct bin for every placed item)"],
           ["Cycle time", "~0.8 s per item"],
           ["Failure handling", "failed items auto-cleared to protect the next pick"],
       ]})
    a({"type": "p", "text":
       "The remaining ~6% are occasional grasp misses, which is realistic for a physical system."})
    a({"type": "image", "path": "docs/figures/fig_results.png", "width_in": 6.0,
       "caption": "Figure 6. Per-episode physical placement success and sort accuracy for a "
                  "representative six-episode run of six products each."})

    a({"type": "h1", "text": "14. Limitations and Future Work"})
    a({"type": "p", "text": "Limitations:"})
    a({"type": "bullets", "items": [
        "Simulation only, with no physical robot or sim-to-real validation.",
        "Detection uses the simulator's ground-truth segmentation masks; color is then classified "
        "from real RGB pixels. A learned detector is wired in but not the active path.",
    ]})
    a({"type": "p", "text": "Future work:"})
    a({"type": "bullets", "items": [
        "Train YOLOv8 on synthetic data with domain randomization.",
        "Add visual servoing on the final approach to close the remaining grasp-miss gap.",
        "Grasp quality estimation (antipodal analysis) to reject and re-plan poor grasps.",
        "A realistic depth-sensor noise model.",
        "Dynamic picking off a moving belt (items currently stop at a staging gate).",
    ]})

    a({"type": "h1", "text": "15. Attribution and License"})
    a({"type": "p", "text":
       "The Franka Emika Panda model in assets/franka/ is from the MuJoCo Menagerie by Google "
       "DeepMind, used under its BSD-3-Clause license. Project code is released under the MIT "
       "license."})

    return c


# --------------------------------------------------------------------------- DOCX
def render_docx(content, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = st.add_run(SUBTITLE)
    run.italic = True
    run.font.size = Pt(13)

    for _ in range(3):
        doc.add_paragraph()
    for label in (AUTHOR, DATE, REPO):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.size = Pt(12)
        if label == AUTHOR:
            r.bold = True
            r.font.size = Pt(14)
    doc.add_page_break()

    for block in content:
        bt = block["type"]
        if bt == "h1":
            doc.add_heading(block["text"], level=1)
        elif bt == "h2":
            doc.add_heading(block["text"], level=2)
        elif bt == "p":
            doc.add_paragraph(block["text"])
        elif bt == "bullets":
            for item in block["items"]:
                doc.add_paragraph(item, style="List Bullet")
        elif bt == "code":
            p = doc.add_paragraph()
            r = p.add_run(block["text"])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif bt == "image":
            if Path(block["path"]).exists():
                doc.add_picture(block["path"], width=Inches(block.get("width_in", 6.0)))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(block["caption"])
                cr.italic = True
                cr.font.size = Pt(9)
        elif bt == "table":
            headers = block["headers"]
            rows = block["rows"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            table.autofit = True
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            doc.add_paragraph()

    doc.save(str(path))
    print("wrote", path)


# ---------------------------------------------------------------------------- PDF
def render_pdf(content, path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, PageBreak, ListFlowable, ListItem)
    from reportlab.lib.enums import TA_CENTER

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10.5, leading=15,
                          spaceAfter=6)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=15, spaceBefore=14,
                        spaceAfter=6, textColor=colors.HexColor("#1F3A5F"))
    code = ParagraphStyle("code", parent=styles["Code"], fontName="Courier", fontSize=8.5,
                          leading=11, backColor=colors.HexColor("#F2F2F2"), borderPadding=6,
                          spaceAfter=8, spaceBefore=2)
    title = ParagraphStyle("title", parent=styles["Title"], fontSize=26,
                           textColor=colors.HexColor("#1F3A5F"), alignment=TA_CENTER)
    subtitle = ParagraphStyle("subtitle", parent=styles["Italic"], fontSize=13,
                              alignment=TA_CENTER, spaceAfter=40)
    center = ParagraphStyle("center", parent=body, alignment=TA_CENTER)

    story = [Spacer(1, 5 * cm), Paragraph(TITLE, title), Spacer(1, 0.4 * cm),
             Paragraph(SUBTITLE, subtitle), Spacer(1, 2 * cm),
             Paragraph("<b>%s</b>" % AUTHOR, center), Paragraph(DATE, center),
             Paragraph(REPO, center), PageBreak()]

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    for block in content:
        bt = block["type"]
        if bt in ("h1", "h2"):
            story.append(Paragraph(esc(block["text"]), h1))
        elif bt == "p":
            story.append(Paragraph(esc(block["text"]), body))
        elif bt == "bullets":
            items = [ListItem(Paragraph(esc(i), body), leftIndent=10) for i in block["items"]]
            story.append(ListFlowable(items, bulletType="bullet", start="circle"))
            story.append(Spacer(1, 0.2 * cm))
        elif bt == "code":
            story.append(Paragraph(esc(block["text"]).replace("\n", "<br/>"), code))
        elif bt == "image":
            if Path(block["path"]).exists():
                from reportlab.platypus import Image as RLImage
                from PIL import Image as PILImage
                iw, ih = PILImage.open(block["path"]).size
                w = block.get("width_in", 6.0) * 2.54 * cm
                h = w * ih / iw
                img = RLImage(block["path"], width=w, height=h)
                img.hAlign = "CENTER"
                story.append(img)
                story.append(Paragraph("<i>%s</i>" % esc(block["caption"]), center))
                story.append(Spacer(1, 0.35 * cm))
        elif bt == "table":
            data = [[Paragraph("<b>%s</b>" % esc(h), body) for h in block["headers"]]]
            for row in block["rows"]:
                data.append([Paragraph(esc(str(v)), body) for v in row])
            ncols = len(block["headers"])
            total = 16.5
            if ncols == 2:
                widths = [5.5 * cm, 11.0 * cm]
            elif ncols == 3:
                widths = [4.0 * cm, 2.5 * cm, 10.0 * cm]
            else:
                widths = [total / ncols * cm] * ncols
            tbl = Table(data, colWidths=widths, hAlign="LEFT")
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3A5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F4F6F9")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 0.3 * cm))

    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm, title=TITLE, author=AUTHOR)
    doc.build(story)
    print("wrote", path)


if __name__ == "__main__":
    content = build_content()
    render_docx(content, OUT_DIR / "Vision_Robot_Report.docx")
    render_pdf(content, OUT_DIR / "Vision_Robot_Report.pdf")
